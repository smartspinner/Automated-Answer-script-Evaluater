# -*- coding: utf-8 -*-
"""
Created on Fri Nov 15 05:17:15 2019

@author: Devika
"""

from docx import Document
import os
import csv
answer=[]
answers=[]
points=0
split=0
score=0
options=[]
total=0
toappend=[]
document = Document('BioAnswerKey.docx')
script=Document('BioAnswerScript_Correct.docx')
for p in script.paragraphs:
    answer.append(p.text)
listoflist=[]
for i in range(0,len(answer)):
    new=answer[i].replace('.','')
    new=answer[i].replace(',','')
    new=answer[i].replace('\n','')
    answer[i]=new
    
    if answer[i]!='':
        answers.append(answer[i])
#print(len(listoflist))
#for p in document.paragraphs:
 #   answers.append(p.text)
f=open('BioAnswerKey.csv','r')
spamreader = csv.reader(f, delimiter=',', quotechar='|')
count=0
flag=0
qno=0
#print("ANSWER",answers)
for i in spamreader: #i is the list keywords for an answer
    
    
    #Removing garbage characters
    if(flag==0):
        i[0]=i[0][3:]
        flag+=1
    length=len(i)
    
    #Removing empty keywords
    while(i[length-1]==''):
    
         i.pop()
         length-=1;
    
    #print(i)
    split=3/length;#Mark splitting for each row
    #print(split)
    #print(answers[qno])
    for j in range(0,len(i)):#j is a keyword for the answer
        #print(i[j])
      
        if('/' in i[j]):
            
            options=i[j].split('/')
            #print(options)
            for k in options:
                #print("inside")
                if k in answers[qno]:
                    score+=split
                    #print("added")
                    break
        else:
            #print("here")
            
            if i[j] in answers[qno]:
                score+=split
                #print("added")
                
       
        total+=score
        score=0
    qno+=1
print("Marks awarded:",total)
    
 
      
#for row in spamreader:
 #   print(','.join(row))
